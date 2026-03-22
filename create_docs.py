from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading("API Bisutería - Documentación de Servicios", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Versión: 2.0.0")
doc.add_paragraph("Fecha: 2026-03-22")
doc.add_paragraph()

doc.add_heading("Descripción General", level=1)
doc.add_paragraph(
    "API REST para la gestión de una bisutería. Permite administrar productos y ventas "
    "con validaciones mediante Zod y autenticación JWT."
)

doc.add_heading("Base URL", level=1)
doc.add_paragraph("http://localhost:3000")

doc.add_heading("AUTENTICACIÓN", level=1)
doc.add_paragraph(
    "La API utiliza autenticación JWT (JSON Web Token). "
    "Los endpoints GET (consultar todos y consultar uno) son públicos. "
    "Los endpoints POST, PUT y DELETE requieren autenticación."
)

doc.add_heading("PASOS PARA MODIFICAR DATOS", level=2)
doc.add_paragraph("1. Iniciar sesión con /api/auth/login para obtener un token JWT")
doc.add_paragraph(
    "2. Incluir el token en el header Authorization de las peticiones protegidas"
)
doc.add_paragraph("3. El formato del header es: Authorization: Bearer <token>")
doc.add_paragraph("4. El token expira en 24 horas")

doc.add_heading("Usuarios Predeterminados", level=2)
doc.add_paragraph("admin / admin123")
doc.add_paragraph("vendedor / vendedor123")

doc.add_heading("AUTH - LOGIN", level=1)
doc.add_paragraph("Inicia sesión y devuelve un token JWT.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl -X POST http://localhost:3000/api/auth/login \\")
doc.add_paragraph('  -H "Content-Type: application/json" \\')
doc.add_paragraph('  -d \'{"username":"admin","password":"admin123"}\'')

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("POST http://localhost:3000/api/auth/login")
doc.add_paragraph("Body (raw, JSON):")
doc.add_paragraph('{"username":"admin","password":"admin123"}')

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("200 OK")
doc.add_paragraph(
    '{"token":"eyJhbGciOiJIUzI1NiIs...","user":{"id":1,"username":"admin"}}'
)

p = doc.add_paragraph()
p.add_run("Error (credenciales inválidas):").bold = True
doc.add_paragraph('401: {"error":"Credenciales inválidas"}')

doc.add_heading("AUTH - REGISTER", level=1)
doc.add_paragraph("Registra un nuevo usuario y devuelve un token JWT.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl -X POST http://localhost:3000/api/auth/register \\")
doc.add_paragraph('  -H "Content-Type: application/json" \\')
doc.add_paragraph('  -d \'{"username":"nuevousuario","password":"password123"}\'')

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("POST http://localhost:3000/api/auth/register")
doc.add_paragraph("Body (raw, JSON):")
doc.add_paragraph('{"username":"nuevousuario","password":"password123"}')

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("201 Created")
doc.add_paragraph(
    '{"token":"eyJhbGciOiJIUzI1Ni...","user":{"id":3,"username":"nuevousuario"}}'
)

p = doc.add_paragraph()
p.add_run("Error (usuario ya existe):").bold = True
doc.add_paragraph('409: {"error":"El usuario ya existe"}')

p = doc.add_paragraph()
p.add_run("Error (validación):").bold = True
doc.add_paragraph(
    '400: {"error":[{"path":["username"],"message":"El usuario debe tener al menos 3 caracteres"}]}'
)

doc.add_heading("PRODUCTOS", level=1)

doc.add_heading("GET /api/products (PÚBLICO)", level=2)
doc.add_paragraph("Obtiene todos los productos. No requiere autenticación.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl http://localhost:3000/api/products")

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("GET http://localhost:3000/api/products")

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("200 OK")
doc.add_paragraph(
    '[{"id":1,"name":"Collar de perlas","material":"Plata 925","category":"Collares","price":45000,"stock":15,"minStock":5}]'
)

doc.add_heading("GET /api/products/:id (PÚBLICO)", level=2)
doc.add_paragraph("Obtiene un producto por su ID. No requiere autenticación.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl http://localhost:3000/api/products/1")

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("GET http://localhost:3000/api/products/1")

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("200 OK")
doc.add_paragraph(
    '{"id":1,"name":"Collar de perlas","material":"Plata 925","category":"Collares","price":45000,"stock":15,"minStock":5}'
)

doc.add_heading("POST /api/products (REQUIERE AUTH)", level=2)
doc.add_paragraph("Crea un nuevo producto. Requiere autenticación.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl -X POST http://localhost:3000/api/products \\")
doc.add_paragraph('  -H "Content-Type: application/json" \\')
doc.add_paragraph('  -H "Authorization: Bearer <tu_token>" \\')
doc.add_paragraph(
    '  -d \'{"name":"Collar","material":"Plata","category":"Collares","price":15000,"stock":10,"minStock":5}\''
)

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("POST http://localhost:3000/api/products")
doc.add_paragraph("Headers: Authorization: Bearer <tu_token>")
doc.add_paragraph("Body (raw, JSON):")
doc.add_paragraph(
    '{"name":"Collar","material":"Plata","category":"Collares","price":15000,"stock":10,"minStock":5}'
)

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("201 Created")
doc.add_paragraph(
    '{"id":13,"name":"Collar","material":"Plata","category":"Collares","price":15000,"stock":10,"minStock":5}'
)

p = doc.add_paragraph()
p.add_run("Sin token:").bold = True
doc.add_paragraph('401: {"error":"Token requerido"}')

doc.add_heading("PUT /api/products/:id (REQUIERE AUTH)", level=2)
doc.add_paragraph("Actualiza un producto existente (parcial). Requiere autenticación.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl -X PUT http://localhost:3000/api/products/1 \\")
doc.add_paragraph('  -H "Content-Type: application/json" \\')
doc.add_paragraph('  -H "Authorization: Bearer <tu_token>" \\')
doc.add_paragraph('  -d \'{"name":"Collar actualizado","price":50000}\'')

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("PUT http://localhost:3000/api/products/1")
doc.add_paragraph("Headers: Authorization: Bearer <tu_token>")
doc.add_paragraph("Body (raw, JSON):")
doc.add_paragraph('{"name":"Collar actualizado","price":50000}')

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("200 OK")
doc.add_paragraph(
    '{"id":1,"name":"Collar actualizado","material":"Plata 925","category":"Collares","price":50000,"stock":15,"minStock":5}'
)

doc.add_heading("DELETE /api/products/:id (REQUIERE AUTH)", level=2)
doc.add_paragraph("Elimina un producto por su ID. Requiere autenticación.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl -X DELETE http://localhost:3000/api/products/1 \\")
doc.add_paragraph('  -H "Authorization: Bearer <tu_token>"')

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("DELETE http://localhost:3000/api/products/1")
doc.add_paragraph("Headers: Authorization: Bearer <tu_token>")

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("204 No Content")

doc.add_heading("VENTAS", level=1)

doc.add_heading("GET /api/sales (PÚBLICO)", level=2)
doc.add_paragraph("Obtiene todas las ventas. No requiere autenticación.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl http://localhost:3000/api/sales")

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("GET http://localhost:3000/api/sales")

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("200 OK")
doc.add_paragraph(
    '[{"id":1,"productId":1,"productName":"Collar de perlas","quantity":2,"unitPrice":45000,"total":90000,"date":"2026-03-15"}]'
)

doc.add_heading("GET /api/sales/:id (PÚBLICO)", level=2)
doc.add_paragraph("Obtiene una venta por su ID. No requiere autenticación.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl http://localhost:3000/api/sales/1")

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("GET http://localhost:3000/api/sales/1")

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("200 OK")
doc.add_paragraph(
    '{"id":1,"productId":1,"productName":"Collar de perlas","quantity":2,"unitPrice":45000,"total":90000,"date":"2026-03-15"}'
)

doc.add_heading("POST /api/sales (REQUIERE AUTH)", level=2)
doc.add_paragraph("Crea una nueva venta. Requiere autenticación.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl -X POST http://localhost:3000/api/sales \\")
doc.add_paragraph('  -H "Content-Type: application/json" \\')
doc.add_paragraph('  -H "Authorization: Bearer <tu_token>" \\')
doc.add_paragraph(
    '  -d \'{"productId":1,"productName":"Collar de perlas","quantity":2,"unitPrice":45000,"total":90000,"date":"2026-03-22"}\''
)

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("POST http://localhost:3000/api/sales")
doc.add_paragraph("Headers: Authorization: Bearer <tu_token>")
doc.add_paragraph("Body (raw, JSON):")
doc.add_paragraph(
    '{"productId":1,"productName":"Collar de perlas","quantity":2,"unitPrice":45000,"total":90000,"date":"2026-03-22"}'
)

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("201 Created")
doc.add_paragraph(
    '{"id":4,"productId":1,"productName":"Collar de perlas","quantity":2,"unitPrice":45000,"total":90000,"date":"2026-03-22"}'
)

p = doc.add_paragraph()
p.add_run("Sin token:").bold = True
doc.add_paragraph('401: {"error":"Token requerido"}')

doc.add_heading("DELETE /api/sales/:id (REQUIERE AUTH)", level=2)
doc.add_paragraph("Elimina una venta por su ID. Requiere autenticación.")

p = doc.add_paragraph()
p.add_run("curl:").bold = True
doc.add_paragraph("curl -X DELETE http://localhost:3000/api/sales/1 \\")
doc.add_paragraph('  -H "Authorization: Bearer <tu_token>"')

p = doc.add_paragraph()
p.add_run("Postman (JSON):").bold = True
doc.add_paragraph("DELETE http://localhost:3000/api/sales/1")
doc.add_paragraph("Headers: Authorization: Bearer <tu_token>")

p = doc.add_paragraph()
p.add_run("Respuesta:").bold = True
doc.add_paragraph("204 No Content")

doc.add_heading("CÓDIGOS DE RESPUESTA", level=1)

table_codes = doc.add_table(rows=8, cols=2)
table_codes.style = "Table Grid"
table_codes.rows[0].cells[0].text = "Código"
table_codes.rows[0].cells[1].text = "Descripción"
table_codes.rows[1].cells[0].text = "200 OK"
table_codes.rows[1].cells[1].text = "Solicitud exitosa"
table_codes.rows[2].cells[0].text = "201 Created"
table_codes.rows[2].cells[1].text = "Recurso creado exitosamente"
table_codes.rows[3].cells[0].text = "204 No Content"
table_codes.rows[3].cells[1].text = "Recurso eliminado exitosamente"
table_codes.rows[4].cells[0].text = "400 Bad Request"
table_codes.rows[4].cells[1].text = "Error de validación"
table_codes.rows[5].cells[0].text = "401 Unauthorized"
table_codes.rows[5].cells[1].text = "Token requerido o credenciales inválidas"
table_codes.rows[6].cells[0].text = "403 Forbidden"
table_codes.rows[6].cells[1].text = "Token inválido o expirado"
table_codes.rows[7].cells[0].text = "409 Conflict"
table_codes.rows[7].cells[1].text = "Recurso ya existe"

doc.add_paragraph()

doc.add_heading("ESQUEMAS DE VALIDACIÓN", level=1)

doc.add_heading("User (para login/register)", level=2)
table_user = doc.add_table(rows=3, cols=3)
table_user.style = "Table Grid"
hdr_user = table_user.rows[0].cells
hdr_user[0].text = "Campo"
hdr_user[1].text = "Tipo"
hdr_user[2].text = "Descripción"
user_data = [
    ("username", "string", "Usuario (mínimo 3 caracteres, requerido)"),
    ("password", "string", "Contraseña (mínimo 6 caracteres, requerido)"),
]
for i, (field, tipo, desc) in enumerate(user_data):
    row_cells = table_user.rows[i].cells
    row_cells[0].text = field
    row_cells[1].text = tipo
    row_cells[2].text = desc

doc.add_paragraph()

doc.add_heading("Product", level=2)
table = doc.add_table(rows=8, cols=3)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Campo"
hdr_cells[1].text = "Tipo"
hdr_cells[2].text = "Descripción"
rows_data = [
    ("id", "number", "Identificador único (auto-generado)"),
    ("name", "string", "Nombre del producto (requerido)"),
    ("material", "string", "Material del producto (requerido)"),
    ("category", "string", "Categoría del producto (requerido)"),
    ("price", "number", "Precio (positivo, requerido)"),
    ("stock", "number", "Stock actual (entero no negativo)"),
    ("minStock", "number", "Stock mínimo (entero no negativo)"),
]
for i, (field, tipo, desc) in enumerate(rows_data):
    row_cells = table.rows[i].cells
    row_cells[0].text = field
    row_cells[1].text = tipo
    row_cells[2].text = desc

doc.add_paragraph()

doc.add_heading("Sale", level=2)
table2 = doc.add_table(rows=8, cols=3)
table2.style = "Table Grid"
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = "Campo"
hdr_cells2[1].text = "Tipo"
hdr_cells2[2].text = "Descripción"
sales_data = [
    ("id", "number", "Identificador único (auto-generado)"),
    ("productId", "number", "ID del producto vendido"),
    ("productName", "string", "Nombre del producto (requerido)"),
    ("quantity", "number", "Cantidad vendida (entero positivo)"),
    ("unitPrice", "number", "Precio unitario (positivo)"),
    ("total", "number", "Total de la venta"),
    ("date", "string", "Fecha de la venta (YYYY-MM-DD)"),
]
for i, (field, tipo, desc) in enumerate(sales_data):
    row_cells2 = table2.rows[i].cells
    row_cells2[0].text = field
    row_cells2[1].text = tipo
    row_cells2[2].text = desc

doc.add_heading("EJECUTAR EL SERVIDOR", level=1)
p = doc.add_paragraph()
p.add_run("Desarrollo (con auto-reload):").bold = True
doc.add_paragraph("npm run dev")
p = doc.add_paragraph()
p.add_run("Producción:").bold = True
doc.add_paragraph("npm start")

doc.save("API_Bisuteria_Documentacion_v4.docx")
print("Documento creado: API_Bisuteria_Documentacion_v4.docx")
